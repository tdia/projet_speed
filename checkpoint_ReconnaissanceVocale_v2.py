import streamlit as st
import speech_recognition as sr
from streamlit_mic_recorder import mic_recorder
import io
import time
import google.generativeai as genai
import os
from gtts import gTTS
import io
import base64
import asyncio
import edge_tts
import fitz  # PyMuPDF
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import moviepy as mp

# ================== CONFIGURATION PAGE ==================
st.set_page_config(
    page_title="XAMLE IA - STUDIO",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================== CONFIGURATION GEMINI ==================
# Note: Idéalement, utilisez st.secrets ou une variable d'environnement
# Pour ce checkpoint, on utilise la clé fournie
GEMINI_API_KEY = "AIzaSyDVgHAvnclF_ethXsNNQCA_HGVtr-4HlNk"
genai.configure(api_key=GEMINI_API_KEY)

# ================== DESIGN / UI (CSS) ==================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&display=swap');
    * { font-family: 'Outfit', sans-serif !important; box-sizing: border-box; }
    /* FOND PRINCIPAL : slate clair premium */
    .stApp { background: linear-gradient(160deg, #f0f4ff 0%, #e8f0fe 50%, #f5f3ff 100%) !important; }
    .block-container { padding: 0 !important; max-width: 100% !important; }
    .stMainBlockContainer { padding: 0 1.5rem 2rem 1.5rem !important; }

    /* LOGIN CARD */
    div[data-testid="stForm"] {
        background: rgba(255,255,255,0.85) !important;
        backdrop-filter: blur(24px) !important;
        -webkit-backdrop-filter: blur(24px) !important;
        border: 1px solid rgba(99,102,241,0.15) !important;
        border-radius: 28px !important;
        padding: 48px 44px 40px !important;
        width: 100% !important;
        max-width: 440px !important;
        box-shadow: 0 20px 60px rgba(99,102,241,0.12), 0 4px 20px rgba(0,0,0,0.06) !important;
        margin: 0 auto !important;
        text-align: center !important;
    }
    div[data-testid="stForm"] > div { border: none !important; }

    /* Logo centré dans le formulaire login */
    div[data-testid="stForm"] [data-testid="stImage"] {
        display: flex !important; justify-content: center !important;
        margin: 0 auto 12px auto !important;
        animation: float 3s ease-in-out infinite;
        filter: drop-shadow(0 4px 16px rgba(99,102,241,0.25));
    }
    div[data-testid="stForm"] [data-testid="stImage"] img {
        border-radius: 20px !important;
        max-width: 110px !important;
    }
    @keyframes float { 0%,100%{transform:translateY(0);} 50%{transform:translateY(-8px);} }
    .login-title {
        background: linear-gradient(90deg, #4f46e5, #7c3aed, #9333ea);
        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
        font-size: 2.4rem; font-weight: 800; margin-bottom: 6px; letter-spacing: -1.5px;
        text-align: center;
    }
    .login-subtitle { color: #64748b; font-size: 0.95rem; margin-bottom: 32px; line-height: 1.5; text-align: center; }

    /* Logo sidebar centré */
    section[data-testid="stSidebar"] [data-testid="stImage"] {
        display: flex !important; justify-content: center !important;
        margin: 12px auto 4px auto !important;
    }
    section[data-testid="stSidebar"] [data-testid="stImage"] img {
        border-radius: 16px !important;
    }

    /* BOUTONS */
    .stButton button {
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%) !important;
        color: white !important; border: none !important;
        border-radius: 14px !important; font-weight: 700 !important;
        font-size: 1rem !important; height: 52px !important;
        box-shadow: 0 6px 18px rgba(79,70,229,0.3) !important;
        transition: all 0.25s ease !important; letter-spacing: 0.3px !important;
    }
    .stButton button:hover { transform: translateY(-2px) !important; box-shadow: 0 10px 24px rgba(79,70,229,0.4) !important; }
    .stButton button[kind="secondary"] {
        background: white !important;
        border: 1.5px solid #e2e8f0 !important;
        color: #64748b !important; box-shadow: 0 1px 4px rgba(0,0,0,0.06) !important;
    }
    .stButton button[kind="secondary"]:hover { border-color: #c7d2fe !important; color: #4f46e5 !important; }
    /* INPUTS login */
    div[data-testid="stTextInput"] label p {
        color: #6366f1 !important; font-weight: 600 !important;
        text-align: left !important; font-size: 0.82rem !important;
        text-transform: uppercase !important; letter-spacing: 0.5px !important;
    }
    div[data-testid="stTextInput"] input {
        background: #f8faff !important;
        border: 1.5px solid #e0e7ff !important;
        color: #1e293b !important; border-radius: 14px !important; height: 52px !important;
    }
    div[data-testid="stTextInput"] input:focus {
        border-color: #6366f1 !important;
        box-shadow: 0 0 0 3px rgba(99,102,241,0.12) !important;
    }

    /* NAV BAR — Pill blanc clean */
    .nav-container {
        background: rgba(255,255,255,0.9) !important;
        backdrop-filter: blur(16px) !important;
        border-radius: 100px !important; padding: 6px 20px !important;
        box-shadow: 0 4px 20px rgba(99,102,241,0.1), 0 1px 3px rgba(0,0,0,0.06) !important;
        margin: 1.5rem auto 2.5rem auto !important; max-width: 96% !important;
        border: 1px solid rgba(99,102,241,0.1) !important;
    }
    div[data-testid="stRadio"] > div { gap: 4px !important; }
    div[data-testid="stRadio"] label {
        background: transparent !important; border: none !important;
        color: #64748b !important; font-weight: 600 !important;
        font-size: 0.9rem !important; padding: 8px 18px !important;
        border-radius: 50px !important; transition: all 0.25s ease !important;
    }
    div[data-testid="stRadio"] label:hover { color: #4f46e5 !important; background: rgba(99,102,241,0.08) !important; }
    div[data-testid="stRadio"] label[data-checked="true"] {
        color: white !important;
        background: linear-gradient(135deg, #4f46e5, #7c3aed) !important;
        box-shadow: 0 4px 12px rgba(79,70,229,0.25) !important;
    }
    div[data-testid="stRadio"] [data-testid="stWidgetLabel"],
    div[data-testid="stRadio"] div[role="radiogroup"] > label > div:first-child { display: none !important; }
    div[data-testid="stRadio"] span[data-testid="stMarkdownContainer"] p { margin: 0 !important; }
    div[data-testid="stRadio"] [data-testid="stMarkdownContainer"] { font-size: 0.9rem !important; }

    /* SIDEBAR */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #f8faff 0%, #f3f0ff 100%) !important;
        border-right: 1px solid #e0e7ff !important;
    }
    section[data-testid="stSidebar"] hr { border-color: #e0e7ff !important; }
    section[data-testid="stSidebar"] label p { color: #475569 !important; }
    section[data-testid="stSidebar"] .stSubheader, section[data-testid="stSidebar"] h3 { color: #1e293b !important; }
    /* TEXTAREA */
    textarea {
        background: #ffffff !important;
        border: 1.5px solid #e0e7ff !important;
        color: #1e293b !important; border-radius: 14px !important;
        font-size: 0.95rem !important; line-height: 1.7 !important;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04) !important;
    }
    textarea:focus { border-color: #6366f1 !important; box-shadow: 0 0 0 3px rgba(99,102,241,0.1) !important; }
    /* FILE UPLOADER */
    [data-testid="stFileUploader"] {
        background: rgba(99,102,241,0.03) !important;
        border: 2px dashed rgba(99,102,241,0.3) !important;
        border-radius: 18px !important;
    }
    [data-testid="stFileUploader"]:hover { border-color: #6366f1 !important; background: rgba(99,102,241,0.06) !important; }
    /* EXPANDER */
    details { background: white !important; border: 1px solid #e0e7ff !important; border-radius: 14px !important; box-shadow: 0 1px 4px rgba(0,0,0,0.04) !important; }
    details summary { color: #475569 !important; font-weight: 600 !important; }
    /* TYPOGRAPHY */
    h1,h2,h3,h4 { color: #1e293b !important; } p,li { color: #475569; }
    hr { border-color: #e2e8f0 !important; margin: 1.5rem 0 !important; }
    /* DOWNLOAD BTN */
    [data-testid="stDownloadButton"] button {
        background: rgba(16,185,129,0.08) !important;
        border: 1.5px solid rgba(16,185,129,0.3) !important;
        color: #059669 !important; box-shadow: none !important;
    }
    ::-webkit-scrollbar { width: 6px; }
    ::-webkit-scrollbar-track { background: #f1f5f9; }
    ::-webkit-scrollbar-thumb { background: #c7d2fe; border-radius: 50px; }
    ::-webkit-scrollbar-thumb:hover { background: #a5b4fc; }
    /* GOOGLE BTN */
    .google-btn {
        display: flex; align-items: center; justify-content: center; gap: 12px;
        background: #f8faff !important; color: #1e293b !important;
        border: 1.5px solid #e0e7ff !important; border-radius: 14px !important;
        height: 52px; margin-top: 16px; font-weight: 600; cursor: pointer;
        text-decoration: none; font-size: 0.95rem; transition: all 0.2s ease;
        box-shadow: 0 1px 4px rgba(0,0,0,0.06);
    }
    .google-btn:hover { background: #f0f4ff !important; border-color: #c7d2fe !important; }
</style>
""", unsafe_allow_html=True)

# ================== LOGIC / UTILS ==================

import tempfile
import requests

def generate_google_cloud_tts(text, voice_name, api_key):
    """Génère de l'audio via l'API Google Cloud TTS (Premium)."""
    url = f"https://texttospeech.googleapis.com/v1/text:synthesize?key={api_key}"
    
    # Détecter la langue à partir du nom de la voix (ex: fr-FR-Studio-A)
    lang_code = "-".join(voice_name.split("-")[:2])
    
    payload = {
        "input": {"text": text},
        "voice": {"languageCode": lang_code, "name": voice_name},
        "audioConfig": {"audioEncoding": "MP3", "pitch": -2.0, "speakingRate": 0.9}
    }
    
    response = requests.post(url, json=payload)
    if response.status_code == 200:
        return base64.b64decode(response.json()["audioContent"])
    else:
        raise Exception(f"Erreur Google Cloud TTS: {response.text}")

# Modèles confirmés disponibles (vérifiés via genai.list_models())
GEMINI_MODELS_PRIORITY = [
    "gemini-2.5-flash-lite",    # Ultra-rapide, idéal pour la traduction
    "gemini-2.0-flash-lite",    # Rapide et léger
    "gemini-2.0-flash",         # Standard fiable
    "gemini-2.5-flash",         # Plus puissant si besoin
    "gemini-flash-latest",      # Alias vers le flash le plus récent
]

# Modèles prioritaires pour la traduction Wolof
GEMINI_WOLOF_MODELS = [
    "gemini-2.5-flash",         # Meilleur équilibre qualité/vitesse pour le Wolof
    "gemini-2.0-flash",         # Fallback fiable
    "gemini-2.5-flash-lite",    # Dernier recours rapide
]

def generate_gemini_ai_audio(text, target_lang):
    """Génère de l'audio via l'IA Gemini avec basculement dynamique entre les modèles."""
    last_error = ""
    for model_name in GEMINI_MODELS_PRIORITY:
        try:
            model_audio = genai.GenerativeModel(model_name)
            
            prompt = f"""
            RÔLE : Tu es Charon, une voix-off de studio masculine, profonde et calme.
            MISSION : Prononce cette phrase en {target_lang} avec une fluidité absolue.
            Marque des pauses légères aux virgules. Ton naturel et professionnel.
            Phrase : {text}
            """
            
            response = model_audio.generate_content(
                prompt,
                generation_config={"response_mime_type": "audio/wav"}
            )
            
            # Extraire les bytes audio
            for part in response.candidates[0].content.parts:
                if part.inline_data.mime_type == "audio/wav":
                    return part.inline_data.data
        except Exception as e:
            last_error = str(e)
            if "429" in last_error or "quota" in last_error.lower() or "404" in last_error:
                continue # Essayer le modèle suivant
            else:
                break
    
    raise Exception(f"Tous les modèles Gemini Audio sont épuisés ou indisponibles. Dernière erreur: {last_error}")

def transcribe_audio_bytes(audio_bytes, language="fr-FR"):
    """Transcription avec basculement automatique entre les modèles Gemini."""
    last_error = ""
    
    # Créer un fichier temporaire unique
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
        tmp_file.write(audio_bytes)
        tmp_path = tmp_file.name

    try:
        uploaded_audio = genai.upload_file(path=tmp_path)
        language_name = {"fr-FR": "Français", "wo-SN": "Wolof", "en-US": "Anglais", "es-ES": "Espagnol", "de-DE": "Allemand", "it-IT": "Italien"}.get(language, language)
        
        prompt = f"""
        MISSION : Réalise une transcription ultra-fidèle de cet audio en {language_name}.
        IMPORTANT :
        1. Capture l'EXPRESSIVITÉ : Indique entre parenthèses les tons (ex: [ton joyeux], [ton sérieux], [pause longue]).
        2. RESPECTE les silences : Utilise des virgules et des points pour refléter le rythme original.
        3. TEXTE PUR : Ne donne que la transcription, sans introduction.
        """

        for model_name in GEMINI_MODELS_PRIORITY:
            try:
                model = genai.GenerativeModel(model_name)
                response = model.generate_content([uploaded_audio, prompt])
                return response.text.strip()
            except Exception as e:
                last_error = str(e)
                if "429" in last_error or "quota" in last_error.lower() or "404" in last_error:
                    continue # Passer au modèle suivant
                else:
                    break
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    return f"❌ Erreur de transcription (Tous modèles épuisés) : {last_error}"

def _detect_and_fix_repetition(text: str) -> str:
    """Détecte et coupe les boucles de répétition dans le texte généré."""
    import re
    # Détecter un mot répété plus de 5 fois de suite (ex: "keneen keneen keneen...")
    pattern = re.compile(r'\b(\w+)\s+(\1\s+){4,}', re.IGNORECASE)
    match = pattern.search(text)
    if match:
        # Tronquer au début de la répétition
        text = text[:match.start()].rstrip() + "..."
    return text


def _build_wolof_prompt(paragraph: str) -> str:
    """Construit le prompt Wolof avec vocabulaire guidé pour un sens clair et naturel."""
    return (
        "Tu es un traducteur expert Francais vers Wolof, originaire de Dakar.\n"
        "Tu parles le Wolof courant des entrepreneurs, des commercants et des animateurs radio senegalais.\n"
        "\n"
        "MISSION : Traduire le texte en Wolof qui a du SENS, est FLUIDE et COMPREHENSIBLE\n"
        "par n'importe quel Senegalais, meme sans education formelle.\n"
        "\n"
        "FORMAT : chaque phrase wolof suivie de sa traduction francaise entre parentheses.\n"
        "  Phrase Wolof naturelle. (Traduction francaise.)\n"
        "Paragraphes separes par une ligne vide.\n"
        "\n"
        "VOCABULAIRE WOLOF COURANT A UTILISER (phonetique A-Z) :\n"
        "  jaay = vendre | jaaykat = vendeur | kilyan = client | accueil = accueil\n"
        "  deglou = ecouter | wakh = parler/dire | sonou = sourire | sourire = sourire\n"
        "  bakh = bon/bien | rafet = beau/belle | nekhte = agreable/content\n"
        "  dieuredieuf = merci | mangi dem = je pars | mangi fi = je suis ici\n"
        "  xam = savoir | xam xam = connaissance/savoir | bokk = partager\n"
        "  tekki = reussir/resoudre | mbootaay = equipe/groupe\n"
        "  foofu = la-bas | fii = ici | sopp = aimer | soxor = orgueil | teguine = respect\n"
        "  gueun = mieux/davantage | yem = suffire | am = avoir | dem = aller\n"
        "  dieul = donner | jot = recevoir/obtenir | dalal = accueillir\n"
        "  wollou = confiance/faire confiance | deeg = comprendre/entendre\n"
        "  rew mi = le pays | dakar = Dakar | yalla = Dieu | ndam = chance/benediction\n"
        "\n"
        "EXEMPLES DE BONNE TRADUCTION :\n"
        "  Nio ngui lene dalal ak jaamu. (Bienvenue et soyez les bienvenus.)\n"
        "  Jaaykat bou bakh man naa deglou kilyan bi te wone ko li ko diakhale. (Un bon vendeur sait ecouter le client et lui montrer ce dont il a besoin.)\n"
        "  Sonou bi moo takh kilyan bi nekh ci sa jaay. (C'est le sourire qui rend le client heureux dans ta boutique.)\n"
        "  Bou nga wakh deug, kilyan bi dina delloussi ak rafet. (Si tu es honnete, le client reviendra avec plaisir.)\n"
        "  Xam xam moo am solo, du xaliss rekk. (Le savoir a de la valeur, pas seulement l'argent.)\n"
        "  Dieuredieuf, yalna Yalla may leen ndam ci liggey bi. (Merci, que Dieu vous benisse dans votre travail.)\n"
        "\n"
        "REGLES ABSOLUES :\n"
        "1. Lettres A-Z et accents francais (e, e, a) UNIQUEMENT. Pas de caracteres speciaux wolof.\n"
        "2. Chaque phrase wolof doit avoir un VRAI SENS — pas une transliteration mot a mot.\n"
        "3. INTERDICTION de repeter le meme mot plus de 2 fois dans une phrase.\n"
        "4. Garde en francais : noms propres, termes techniques (boutique, client, produit, pagne, CFA).\n"
        "5. Adapte les expressions : traduis le SENS et l'EMOTION, pas les mots litteralement.\n"
        "\n"
        f"Texte francais a traduire :\n{paragraph}"
    )


def translate_text(text, target_language):
    """Traduction par paragraphes avec anti-boucle et cache session."""
    import re

    # ---- Cache ----
    cache_key = f"_cache_{target_language}_{hash(text[:200])}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]

    # ---- Pré-traitement : supprimer annotations [ton X] [pause X] ----
    text_clean = re.sub(r'\[ton\s+[^\]]+\]', '', text)
    text_clean = re.sub(r'\[pause\s+[^\]]+\]', ', ', text_clean)
    text_clean = re.sub(r'\[pause\]', ', ', text_clean)
    text_clean = re.sub(r'  +', ' ', text_clean).strip()

    # ---- Pour les langues autres que Wolof : traduction directe ----
    if "Wolof" not in target_language:
        prompt = f"Traduis ce texte en {target_language}. Reponds uniquement avec la traduction :\n\n{text_clean[:6000]}"
        errors = []
        for model_name in GEMINI_MODELS_PRIORITY:
            try:
                model = genai.GenerativeModel(model_name,
                    generation_config={"max_output_tokens": 2048, "temperature": 0.5})
                result = model.generate_content(prompt).text.strip()
                st.session_state[cache_key] = result
                return result
            except Exception as e:
                errors.append(f"{model_name}: {str(e)[:50]}")
        return "Traduction echouee : " + " | ".join(errors)

    # ---- Wolof : découpage en paragraphes naturels (anti-boucle) ----
    raw_paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text_clean) if p.strip()]

    # Si un seul bloc sans paragraphes, découper par sections (1. 2. 3.) ou sauts de ligne
    if len(raw_paragraphs) == 1 and len(raw_paragraphs[0]) > 800:
        sentences = re.split(r'(?<=[.!?»])\s+(?=\d+\.)|(?<=[.!?»])\s+', raw_paragraphs[0])
        chunks, current = [], ""
        for s in sentences:
            if len(current) + len(s) < 1200:
                current += " " + s
            else:
                if current:
                    chunks.append(current.strip())
                current = s
        if current:
            chunks.append(current.strip())
        raw_paragraphs = [c for c in chunks if c] if chunks else raw_paragraphs

    translated_parts = []

    for paragraph in raw_paragraphs:
        if not paragraph:
            continue

        prompt = _build_wolof_prompt(paragraph)
        errors = []
        part_result = None

        for model_name in GEMINI_WOLOF_MODELS:
            try:
                model = genai.GenerativeModel(
                    model_name,
                    generation_config={
                        "max_output_tokens": 8192,  # Pas de limite artificielle
                        "temperature": 0.6,
                        "candidate_count": 1,
                    }
                )
                response = model.generate_content(prompt)
                part = response.text.strip()

                # Anti-boucle : détecter et couper les répétitions
                part = _detect_and_fix_repetition(part)
                part = _clean_wolof_chars(part)

                part_result = part
                break
            except Exception as e:
                errors.append(f"{model_name}: {str(e)[:40]}")
                continue

        if part_result:
            translated_parts.append(part_result)
        elif errors:
            translated_parts.append(f"[Erreur paragraphe : {errors[0]}]")

    # ---- Assemblage final ----
    result = "\n\n".join(translated_parts)
    st.session_state[cache_key] = result
    return result


def _clean_wolof_chars(text: str) -> str:
    """Remplace les caracteres speciaux Wolof par leurs equivalents phonetiques A-Z."""
    subs = [
        ("\u00eb", "eu"), ("\u00cb", "Eu"),   # e ouvert wolof -> eu
        ("\u00f1", "gn"), ("\u00d1", "Gn"),   # n tilde -> gn
        ("\u014b", "ng"), ("\u014a", "Ng"),   # eng -> ng
        ("\u0253", "b"),                        # b implosif -> b
        ("\u0257", "d"),                        # d implosif -> d
        ("\u02bc", "'"),  ("\u2019", "'"),     # apostrophes typographiques -> '
        ("\u00f3", "o"),  ("\u00d3", "O"),     # o accent -> o
        # Le x wolof (son kh) est gere par le prompt directement
    ]
    for old, new in subs:
        text = text.replace(old, new)
    return text


# ================== INTERFACE ET AUTHENTIFICATION ==================

def check_password():
    """Gère l'affichage du portail de connexion et valide l'utilisateur."""
    def password_entered():
        # Identifiants de test (Peut être relié à une base de données)
        if st.session_state["username_input"].strip().lower() == "admin" and st.session_state["password_input"] == "admin":
            st.session_state["password_correct"] = True
            st.session_state["logged_user"] = st.session_state["username_input"]
            del st.session_state["password_input"]  # Sécurité: Ne pas garder le mot de passe en mémoire
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # Centrage avec des colonnes Streamlit (plus fiable)
        _, center_col, _ = st.columns([1, 1.2, 1])
        with center_col:
            st.markdown('<div style="height: 10vh;"></div>', unsafe_allow_html=True)
            with st.form("login_form"):
                # Logo centré en base64
                try:
                    with open("pos_logo.png.jpg", "rb") as _f:
                        _logo_b64 = base64.b64encode(_f.read()).decode()
                    st.markdown(f'<div style="text-align:center;margin-bottom:12px;"><img src="data:image/jpeg;base64,{_logo_b64}" style="width:110px;border-radius:18px;filter:drop-shadow(0 4px 16px rgba(99,102,241,0.3));animation:float 3s ease-in-out infinite;"></div>', unsafe_allow_html=True)
                except:
                    st.markdown('<div style="text-align:center;font-size:3rem;margin-bottom:8px;">&#10019;</div>', unsafe_allow_html=True)
                st.markdown('<div class="login-title">Xamle Studio AI</div>', unsafe_allow_html=True)
                st.markdown('<div class="login-subtitle">Connectez-vous pour accéder au Studio Business</div>', unsafe_allow_html=True)
                
                st.text_input("Identifiant", key="username_input", placeholder="admin")
                st.text_input("Mot de passe", type="password", key="password_input", placeholder="••••••••")
                st.form_submit_button("Continuer →", on_click=password_entered, use_container_width=True)
                
                st.markdown("""
                    <a href="#" class="google-btn">
                        <img src="https://www.gstatic.com/images/branding/product/1x/gsa_512dp.png" width="20">
                        Se connecter avec Google
                    </a>
                """, unsafe_allow_html=True)
        return False
        
    elif not st.session_state["password_correct"]:
        _, center_col, _ = st.columns([1, 1.2, 1])
        with center_col:
            st.markdown('<div style="height: 10vh;"></div>', unsafe_allow_html=True)
            with st.form("login_form_err"):
                # Logo centré en base64
                try:
                    with open("pos_logo.png.jpg", "rb") as _f:
                        _logo_b64 = base64.b64encode(_f.read()).decode()
                    st.markdown(f'<div style="text-align:center;margin-bottom:12px;"><img src="data:image/jpeg;base64,{_logo_b64}" style="width:110px;border-radius:18px;filter:drop-shadow(0 4px 16px rgba(99,102,241,0.3));"></div>', unsafe_allow_html=True)
                except:
                    st.markdown('<div style="text-align:center;font-size:3rem;margin-bottom:8px;">&#10019;</div>', unsafe_allow_html=True)
                st.markdown('<div class="login-title">Xamle Studio AI</div>', unsafe_allow_html=True)
                st.markdown('<div class="login-subtitle" style="color:#ef4444;">Identifiants incorrects. Réessayez.</div>', unsafe_allow_html=True)
                
                st.text_input("Identifiant", key="username_input")
                st.text_input("Mot de passe", type="password", key="password_input")
                st.form_submit_button("Réessayer →", on_click=password_entered, use_container_width=True)
        return False
    else:
        # Utilisateur validé
        return True

def main():
    if not check_password():
        return
        
    # NAVIGATION BAR (MagicPattern Style)
    st.markdown('<div class="nav-container">', unsafe_allow_html=True)
    n1, n2, n3 = st.columns([1, 4.5, 1])
    
    with n1:
        st.markdown('<div style="color: #1e293b; font-weight: 800; font-size: 1.3rem; letter-spacing: -1px; padding-top: 5px; padding-left:10px;">XAMLE STUDIO</div>', unsafe_allow_html=True)

    with n2:
        selected_menu = st.radio("Menu", ["🎬 Vidéo", "📄 Document", "📁 Charger", "🔴 Enregistrer", "⚡ En direct"], horizontal=True, label_visibility="collapsed", key="main_menu")

    with n3:
        if st.button("Quitter 🚪", key="exit_btn", use_container_width=True):
            st.session_state.clear()
            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    with st.sidebar:
        # Logo centré en base64 dans la sidebar
        try:
            with open("pos_logo.png.jpg", "rb") as _f:
                _logo_b64 = base64.b64encode(_f.read()).decode()
            st.markdown(f'<div style="text-align:center;margin:16px 0 8px 0;"><img src="data:image/jpeg;base64,{_logo_b64}" style="width:120px;border-radius:16px;"></div>', unsafe_allow_html=True)
        except:
            pass
        st.markdown('<div class="login-title" style="font-size:1.4rem;text-align:center;margin-bottom:4px;">Xamle Studio AI</div>', unsafe_allow_html=True)
        st.title("🛡️ Configuration")
        
        src_lang = st.selectbox(
            "Langue source (Audio)",
            ["fr-FR", "wo-SN", "en-US", "es-ES", "de-DE", "it-IT"],
            format_func=lambda x: {"fr-FR": "🇫🇷 Français", "wo-SN": "🇸🇳 Wolof", "en-US": "🇺🇸 Anglais", "es-ES": "🇪🇸 Espagnol", "de-DE": "🇩🇪 Allemand", "it-IT": "🇮🇹 Italien"}.get(x, x)
        )

        st.divider()
        st.subheader("⚙️ Moteur Vocal")
        tts_engine = st.radio("Technologie TTS", ["Edge-TTS (Gratuit)", "Google Cloud (Premium)", "Gemini AI (Voix Native)"], index=2)
        
        st.divider()
        st.subheader("Traduire en...")
        target_lang_name = st.selectbox(
            "Langue cible",
            ["Wolof", "Français", "Anglais", "Arabe", "Espagnol", "Chinois", "Allemand", "Italien"]
        )
        
        # --- SÉLECTION DE LA VOIX ---
        st.divider()
        st.subheader("🔊 Style de Voix")
        
        # Mapping des voix par langue et moteur
        if tts_engine == "Edge-TTS (Gratuit)":
            voice_map = {
                "Wolof": [
                    ("fr-FR-HenriNeural", "👨 Henri (Radio - Recommandé)"),
                    ("fr-FR-RemyMultilingualNeural", "👨 Rémy (Narrateur)"),
                    ("fr-FR-VivienneMultilingualNeural", "👩 Vivienne (Douce)"),
                    ("fr-FR-DeniseNeural", "👩 Denise (Claire)")
                ],
                "Français": [
                    ("fr-FR-HenriNeural", "👨 Henri (Classique)"),
                    ("fr-FR-RemyMultilingualNeural", "👨 Rémy (Pro)"),
                    ("fr-FR-VivienneMultilingualNeural", "👩 Vivienne (Naturelle)")
                ],
                "Anglais": [
                    ("en-US-GuyNeural", "👨 Guy (News)"),
                    ("en-US-AvaNeural", "👩 Ava")
                ]
            }
        elif tts_engine == "Google Cloud (Premium)":
            voice_map = {
                "Wolof": [
                    ("en-US-Journey-F", "🎙️ Charon (Journey Premium - Par défaut)"),
                    ("fr-FR-Studio-D", "👨 Studio Premium (Profond)"),
                    ("fr-FR-Studio-A", "👩 Studio Premium (Voix-off)"),
                    ("fr-FR-Neural2-B", "👨 Neural (Masculin)")
                ],
                "Français": [
                    ("fr-FR-Studio-D", "👨 Studio Male (Qualité CD)"),
                    ("fr-FR-Studio-A", "👩 Studio Female (Qualité CD)"),
                    ("fr-FR-Polyglot-1", "🌍 Voix Multilingue")
                ],
                "Anglais": [
                    ("en-US-Journey-F", "🎙️ Charon (Journey Premium)"),
                    ("en-US-Studio-O", "👨 Studio O (Luxe)"),
                    ("en-US-Journey-D", "👨 Journey Male")
                ]
            }
        else: # Gemini AI Native
             voice_map = {
                "Wolof": [("gemini-native", "🎙️ Gemini Expressif (Direct AI)")],
                "Français": [("gemini-native", "🎙️ Gemini Expressif (Direct AI)")],
                "Anglais": [("gemini-native", "🎙️ Gemini Expressif (Direct AI)")]
            }
        
        default_voices = [("en-US-Journey-F", "🎙️ Charon")] if tts_engine == "Google Cloud (Premium)" else [("fr-FR-HenriNeural", "👨 Henri")]
        if tts_engine == "Gemini AI (Voix Native)": default_voices = [("gemini-native", "🎙️ Gemini Direct")]
        available_voices = voice_map.get(target_lang_name, default_voices)
        
        selected_voice_id = st.selectbox(
            "Choisir la voix spécifique",
            options=[v[0] for v in available_voices],
            format_func=lambda x: next(v[1] for v in available_voices if v[0] == x)
        )
        
        st.divider()
        if st.button("➕ Nouvelle Session", use_container_width=True):
            # Conserver la connexion, mais effacer les données
            user = st.session_state.get("logged_user", "admin")
            pass_ok = st.session_state.get("password_correct", True)
            st.session_state.clear()
            st.session_state["logged_user"] = user
            st.session_state["password_correct"] = pass_ok
            st.rerun()

        if st.button("🚪 Déconnexion", use_container_width=True, type="secondary"):
            st.session_state.clear()
            st.rerun()

        st.divider()
        st.caption("Propulsé par Google Speech & Gemini AI")

    col1, col2 = st.columns([1, 1.2], gap="large")

    with col1:
        # En-tête dynamique contextuel
        page_headers = {
            "🔴 Enregistrer": ("Studio d'Enregistrement", "Capturez votre voix en temps réel", "🔴"),
            "📁 Charger":     ("Importer un Audio", "Analysez vos fichiers audio locaux", "📁"),
            "📄 Document":    ("Traduction de Documents", "PDF vers votre langue cible", "📄"),
            "🎬 Vidéo":       ("Analyse Vidéo", "Extrayez et transcrivez l'audio d'une vidéo", "🎬"),
            "⚡ En direct":   ("Mode Instantané", "Transcription et traduction en direct", "⚡"),
        }
        title, subtitle, icon = page_headers.get(selected_menu, ("Studio", "", "🎙️"))
        st.markdown(f'''
            <div style="margin-bottom:24px;">
                <div style="font-size:2rem;font-weight:800;color:#f1f5f9;letter-spacing:-1px;">{title}</div>
                <div style="color:#64748b;font-size:0.95rem;margin-top:4px;">{subtitle}</div>
            </div>
        ''', unsafe_allow_html=True)

        audio_input = None

        if selected_menu == "🔴 Enregistrer":
            st.markdown('<p style="color:#94a3b8;font-size:0.9rem;margin-bottom:8px;">Appuyez sur le bouton et parlez — l\'IA transcrit automatiquement.</p>', unsafe_allow_html=True)
            audio_data = mic_recorder(start_prompt="🎤 Démarrer l'enregistrement", stop_prompt="⏹️ Arrêter", key='recorder')
            if audio_data:
                st.audio(audio_data['bytes'])
                audio_input = audio_data['bytes']
        
        elif selected_menu == "📁 Charger":
            st.markdown('<p style="color:#94a3b8;font-size:0.9rem;margin-bottom:8px;">Glissez un fichier WAV ou MP3 pour le transcrire et le traduire.</p>', unsafe_allow_html=True)
            up = st.file_uploader("Importer un fichier audio", type=["wav", "mp3"], label_visibility="collapsed")
            if up:
                st.audio(up)
                audio_input = up.read()

        elif selected_menu == "📄 Document":
            st.markdown('<p style="color:#94a3b8;font-size:0.9rem;margin-bottom:12px;">Importez un PDF et obtenez une traduction complète et téléchargeable.</p>', unsafe_allow_html=True)
            doc_up = st.file_uploader("Sélectionner un fichier PDF", type=["pdf"], label_visibility="collapsed")
            
            if doc_up:
                if 'doc_content' not in st.session_state or st.session_state.get('last_doc_name') != doc_up.name:
                    with st.spinner("Extraction du texte du PDF..."):
                        pdf_data = doc_up.read()
                        doc = fitz.open(stream=pdf_data, filetype="pdf")
                        full_text = ""
                        for page in doc:
                            full_text += page.get_text()
                        st.session_state['doc_content'] = full_text
                        st.session_state['last_doc_name'] = doc_up.name
                        if 'doc_translation' in st.session_state: del st.session_state['doc_translation']

                st.success(f"Texte extrait avec succès ({len(st.session_state['doc_content'])} caractères)")
                with st.expander("Voir le texte original"):
                    st.text_area("Texte source", st.session_state['doc_content'], height=200)

                if st.button(f"✨ Traduire le document en {target_lang_name}", use_container_width=True):
                    with st.spinner("Traduction en cours (Wolof Simple)..."):
                        # Gemini gère bien les longs textes, mais on peut chunker si nécessaire.
                        # Pour un studio, on envoie tout si c'est raisonnable (<30k chars)
                        translation = translate_text(st.session_state['doc_content'][:30000], target_lang_name)
                        st.session_state['doc_translation'] = translation

        elif selected_menu == "🎬 Vidéo":
            st.markdown("#### 🎬 Analyse & Traduction de Vidéo")
            vid_up = st.file_uploader("Importer une vidéo (MP4, MOV, AVI)", type=["mp4", "mov", "avi"])
            
            if vid_up:
                st.video(vid_up)
                if st.button("🚀 Extraire & Transcrire l'audio de la vidéo", use_container_width=True):
                    with st.spinner("Extraction de l'audio et analyse des tons..."):
                        # Enregistrer la vidéo temporairement
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as t_vid:
                            t_vid.write(vid_up.read())
                            t_vid_path = t_vid.name
                        
                        try:
                            # Extraire l'audio
                            video = mp.VideoFileClip(t_vid_path)
                            audio_path = t_vid_path.replace(".mp4", ".wav")
                            video.audio.write_audiofile(audio_path, codec='pcm_s16le')
                            
                            with open(audio_path, "rb") as f:
                                audio_bytes = f.read()
                            
                            # Transcrire avec tons/pauses
                            transcription = transcribe_audio_bytes(audio_bytes, src_lang)
                            st.session_state['transcription'] = transcription
                            
                            # Nettoyage
                            video.close()
                            os.remove(t_vid_path)
                            os.remove(audio_path)
                            
                            st.rerun()
                        except Exception as e:
                            st.error(f"Erreur lors du traitement vidéo : {e}")

        else: # ⚡ En direct
            st.markdown("#### ⚡ Transcription & Traduction Instantanée")
            import streamlit.components.v1 as components
            
            js_target_lang = "wo" if "Wolof" in target_lang_name else "fr" if "Français" in target_lang_name else "en" if "Anglais" in target_lang_name else "ar" if "Arabe" in target_lang_name else "es"
            gtts_lang = "fr" if "Wolof" in target_lang_name else "fr" if "Français" in target_lang_name else "en" if "Anglais" in target_lang_name else "ar" if "Arabe" in target_lang_name else "es" if "Espagnol" in target_lang_name else "zh" if "Chinois" in target_lang_name else "de" if "Allemand" in target_lang_name else "en"
            
            catch_speech_js = f"""
                <style>
                    .mic-btn {{
                        background: linear-gradient(135deg, #6366f1 0%, #a855f7 100%);
                        color: white; border: none; width: 80px; height: 80px; border-radius: 50%;
                        cursor: pointer; font-size: 2.5em; display: flex; align-items: center; justify-content: center;
                        box-shadow: 0 8px 20px rgba(99, 102, 241, 0.4); transition: transform 0.2s ease, box-shadow 0.2s ease;
                        margin: 0 auto;
                    }}
                    .mic-btn:hover {{
                        transform: scale(1.05); box-shadow: 0 12px 25px rgba(99, 102, 241, 0.6);
                    }}
                    .stop-btn {{
                        background: #ef4444; color: white; border: none; width: 80px; height: 80px; border-radius: 50%;
                        cursor: pointer; font-size: 2.5em; display: none; align-items: center; justify-content: center;
                        box-shadow: 0 8px 20px rgba(239, 68, 68, 0.4); margin: 0 auto;
                        animation: pulse 1.5s infinite;
                    }}
                    @keyframes pulse {{
                        0% {{ box-shadow: 0 0 0 0 rgba(239, 68, 68, 0.7); }}
                        70% {{ box-shadow: 0 0 0 20px rgba(239, 68, 68, 0); }}
                        100% {{ box-shadow: 0 0 0 0 rgba(239, 68, 68, 0); }}
                    }}
                </style>
                <div style="background: rgba(255,255,255,0.03); padding: 20px; border-radius: 15px; border: 1px solid rgba(255,255,255,0.1); font-family: 'Outfit', sans-serif;">
                    
                    <button id="start-btn" class="mic-btn" title="Activer le mode Live">🎤</button>
                    <button id="stop-btn" class="stop-btn" title="Arrêter la capture">⏹️</button>
                    
                    <div id="status" style="color: #4facfe; font-size: 0.9em; text-align: center; margin-top: 15px; margin-bottom: 25px; font-weight: 600;">Cliquez sur le micro au centre pour parler.</div>
                    
                    <div style="margin-bottom: 5px; color: #94a3b8; font-size: 0.8em;">Transcription :</div>
                    <div id="live-output" style="background: #ffffff; padding: 15px; border-radius: 8px; font-size: 1.1em; min-height: 50px; color: #000000; font-style: italic; margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">...</div>
                    
                    <div style="margin-bottom: 5px; color: #6366f1; font-size: 0.8em;">Traduction :</div>
                    <div id="live-translation" style="background: #ffffff; padding: 15px; border-radius: 8px; font-size: 1.1em; min-height: 50px; color: #000000; font-weight: 600; line-height: 1.4; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">...</div>
                    
                    <div style="margin-top: 15px; margin-bottom: 5px; display: flex; align-items: center; justify-content: center; background: rgba(0,0,0,0.2); padding: 10px; border-radius: 10px;">
                        <input type="checkbox" id="auto-speak-cb" checked style="margin-right: 10px; width: 18px; height: 18px; cursor: pointer; accent-color: #6366f1;">
                        <label for="auto-speak-cb" style="cursor: pointer; color: #f8fafc; font-size: 0.9em; font-weight: 500;">🔊 Lire automatiquement la traduction en public</label>
                    </div>
                    
                    <button id="save-btn" style="background: #10b981; color: white; border: none; padding: 12px; border-radius: 10px; cursor: pointer; font-weight: 600; width: 100%; margin-top: 15px; transition: transform 0.2s;">💾 Enregistrer vers Résultats</button>
                </div>

                <script>
                    const output = document.getElementById('live-output');
                    const transOutput = document.getElementById('live-translation');
                    const startBtn = document.getElementById('start-btn');
                    const stopBtn = document.getElementById('stop-btn');
                    const saveBtn = document.getElementById('save-btn');
                    const status = document.getElementById('status');
                    
                    let recognition;
                    let fullFinalText = "";
                    let fullTranslationText = "";
                    let isStopped = true;
                    
                    // LECTURE VOCALE NATIVE AVEC PHONÉTIQUE
                    let audioQueue = [];
                    let isPlayingAudio = false;

                    function phonetizeWolof(text) {{
                        const isWolof = "{target_lang_name}".includes("Wolof");
                        if (!isWolof) return text;
                        // Algorithme Avancé de Phonétisation Wolof vers Français (Natif)
                        let t = text;
                        t = t.replace(/e/g, "é").replace(/E/g, "É");
                        t = t.replace(/ë/g, "eu").replace(/Ë/g, "Eu");
                        t = t.replace(/c/g, "thi").replace(/C/g, "Thi");
                        t = t.replace(/j/g, "dji").replace(/J/g, "Dji");
                        t = t.replace(/x/g, "kh").replace(/X/g, "Kh");
                        t = t.replace(/ñ/g, "gn").replace(/Ñ/g, "Gn");
                        t = t.replace(/ŋ/g, "ng").replace(/Ŋ/g, "Ng");
                        t = t.replace(/u/g, "ou").replace(/U/g, "Ou");
                        t = t.replace(/q/g, "k").replace(/Q/g, "K");
                        
                        // Garder le son "G" dur avant les voyelles (ex: ngéén ne doit pas dire njéén)
                        t = t.replace(/gé/g, "gué").replace(/Gé/g, "Gué");
                        t = t.replace(/gi/g, "gui").replace(/Gi/g, "Gui");
                        t = t.replace(/geu/g, "gueu").replace(/Geu/g, "Gueu");
                        
                        // Lettrines longues (Wa -> Oua pour éviter le V français)
                        t = t.replace(/wa/g, "oua").replace(/Wa/g, "Oua");
                        t = t.replace(/wé/g, "oué").replace(/Wé/g, "Oué");
                        t = t.replace(/wi/g, "oui").replace(/Wi/g, "Oui");
                        t = t.replace(/wo/g, "ouo").replace(/Wo/g, "Ouo");
                        t = t.replace(/wou/g, "ou").replace(/Wou/g, "Ou");

                        t = t.replace(/aa/g, "a").replace(/éé/g, "é").replace(/ii/g, "i").replace(/oo/g, "o").replace(/ouou/g, "ou");
                        
                        // Éviter la nasalisation française (ex: "Xam" -> khame au lieu de khan, "man" -> manne au lieu de ment)
                        t = t.replace(/([aeiouéoóòu])([nm])\b/gi, "$1$2e");
                        t = t.replace(/([aeiouéoóòu])([nm])(?=[bcdfghjklmnpqrstvwxz])/gi, "$1$2e");
                        
                        return t;
                    }}

                    function playNextAudio() {{
                        if (audioQueue.length === 0) {{
                            isPlayingAudio = false;
                            return;
                        }}
                        isPlayingAudio = true;
                        
                        const textToSpeak = audioQueue.shift();
                        if (!('speechSynthesis' in window)) {{
                            playNextAudio();
                            return;
                        }}

                        // Appliquer la phonétique si la langue cible est le Wolof
                        const phoneticText = phonetizeWolof(textToSpeak);
                        const utterance = new SpeechSynthesisUtterance(phoneticText);
                        
                        let langTTS = "fr-FR";
                        if ("{target_lang_name}".includes("Anglais")) langTTS = "en-US";
                        else if ("{target_lang_name}".includes("Espagnol")) langTTS = "es-ES";
                        else if ("{target_lang_name}".includes("Arabe")) langTTS = "ar-SA";
                        else if ("{target_lang_name}".includes("Chinois")) langTTS = "zh-CN";
                        else if ("{target_lang_name}".includes("Allemand")) langTTS = "de-DE";
                        
                        const voices = window.speechSynthesis.getVoices();
                        // Recherche ciblée sur des voix d'hommes de haute qualité
                        let bestVoice = voices.find(v => v.lang.includes(langTTS.split('-')[0]) && (v.name.includes("Male") || v.name.includes("Henri") || v.name.includes("Paul") || v.name.includes("Thomas") || v.name.includes("Denise") || v.name.includes("Google")));
                        
                        if (!bestVoice) bestVoice = voices.find(v => v.lang.includes(langTTS.split('-')[0]) && (v.name.includes("Natural") || v.name.includes("Premium")));
                        if (!bestVoice) bestVoice = voices.find(v => v.lang.includes(langTTS.split('-')[0]));
                        
                        if (bestVoice) utterance.voice = bestVoice;
                        else utterance.lang = langTTS;
                        
                        utterance.rate = "{target_lang_name}".includes("Wolof") ? 0.85 : 0.95;
                        utterance.pitch = 0.9; // Un peu plus grave pour le côté "Beau"
                        
                        utterance.onend = () => {{ playNextAudio(); }};
                        utterance.onerror = () => {{ playNextAudio(); }};
                        
                        window.speechSynthesis.speak(utterance);
                    }}

                    async function translateSegment(textSegment) {{
                        if (!textSegment.trim()) return;
                        try {{
                            const url = `https://translate.googleapis.com/translate_a/single?client=gtx&sl=auto&tl={js_target_lang}&dt=t&q=${{encodeURIComponent(textSegment)}}`;
                            const response = await fetch(url);
                            const data = await response.json();
                            if (data && data[0]) {{
                                let translatedSegment = "";
                                data[0].forEach(item => {{ translatedSegment += item[0]; }});
                                
                                fullTranslationText += " " + translatedSegment;
                                transOutput.innerHTML = fullTranslationText;
                                
                                // Placer dans la file d'attente vocale (CLOUD GOOGLE PREMIUM)
                                const autoSpeakElem = document.getElementById('auto-speak-cb');
                                if (autoSpeakElem && autoSpeakElem.checked && translatedSegment.trim().length > 0) {{
                                    audioQueue.push(translatedSegment);
                                    if (!isPlayingAudio) {{
                                        playNextAudio();
                                    }}
                                }}
                            }}
                        }} catch(e) {{ console.error("Erreur Traduction:", e); }}
                    }}

                    if ('webkitSpeechRecognition' in window) {{
                        recognition = new webkitSpeechRecognition();
                        recognition.continuous = true;
                        recognition.interimResults = true;
                        recognition.lang = '{src_lang}';

                        recognition.onstart = () => {{
                            status.innerText = "L'IA vous écoute en continu... (Microphone actif)";
                            startBtn.style.display = 'none';
                            stopBtn.style.display = 'block';
                        }};

                        recognition.onresult = (event) => {{
                            let currentFinal = '';
                            let interim = '';
                            for (let i = event.resultIndex; i < event.results.length; ++i) {{
                                if (event.results[i].isFinal) {{
                                    currentFinal += event.results[i][0].transcript;
                                }} else {{
                                    interim += event.results[i][0].transcript;
                                }}
                            }}
                            
                            if (currentFinal) {{
                                fullFinalText += currentFinal;
                                translateSegment(currentFinal); // Traduire et envoyer
                            }}
                            
                            output.innerHTML = '<strong>' + fullFinalText + '</strong> <span style="color:#64748b;">' + interim + '</span>';
                        }};

                        recognition.onend = () => {{
                            if (!isStopped) {{
                                // Auto-reconnexion si coupure automatique par le navigateur (silence, etc)
                                try {{ recognition.start(); }} catch(e) {{}}
                            }} else {{
                                status.innerText = "Mode Direct arrêté.";
                                startBtn.style.display = 'block';
                                stopBtn.style.display = 'none';
                            }}
                        }};

                        startBtn.onclick = () => {{
                            isStopped = false;
                            try {{ recognition.start(); }} catch(e) {{}}
                            if ('speechSynthesis' in window) {{
                                window.speechSynthesis.cancel();
                                // Petite astuce pour débloquer l'audio synthesis sur l'iframe
                                const unlockUtterance = new SpeechSynthesisUtterance("");
                                unlockUtterance.volume = 0;
                                window.speechSynthesis.speak(unlockUtterance);
                            }}
                        }};
                        stopBtn.onclick = () => {{
                            isStopped = true;
                            recognition.stop();
                        }};
                        saveBtn.onclick = () => {{
                            if (fullFinalText.trim() === "") return;
                            status.innerText = "⏳ Enregistrement...";
                            window.parent.postMessage({{
                                type: 'streamlit:set_widget_value',
                                key: 'auto_sync',
                                value: fullFinalText + " ||| " + fullTranslationText
                            }}, '*');
                            setTimeout(() => {{ status.innerText = "✅ Placé dans les résultats !"; }}, 2000);
                        }};
                    }} else {{
                        status.innerText = "Navigateur non compatible.";
                    }}
                </script>
            """
            components.html(catch_speech_js, height=600)
            
            # Récupération automatique et fusion avec l'historique
            sync_data = st.text_input("Buffer Live (Caché)", key="auto_sync", label_visibility="collapsed")
            if sync_data and " ||| " in sync_data:
                if 'last_sync' not in st.session_state: st.session_state.last_sync = ""
                
                if sync_data != st.session_state.last_sync:
                    transcript_seg, translation_seg = sync_data.split(" ||| ", 1)
                    
                    last_t, last_tr = "", ""
                    if st.session_state.last_sync and " ||| " in st.session_state.last_sync:
                        last_t, last_tr = st.session_state.last_sync.split(" ||| ", 1)
                    
                    # Déduction de la nouveauté pour éviter les doublons sur multi-clic
                    if transcript_seg.startswith(last_t) and translation_seg.startswith(last_tr):
                        new_transcript = transcript_seg[len(last_t):]
                        new_translation = translation_seg[len(last_tr):]
                    else:
                        new_transcript = transcript_seg
                        new_translation = translation_seg
                    
                    if 'transcription' not in st.session_state: st.session_state.transcription = ""
                    if 'translation' not in st.session_state: st.session_state.translation = ""
                    
                    # Accumulation fluide sans effacer l'existant
                    st.session_state.transcription += f" {new_transcript}"
                    st.session_state.translation += f" {new_translation}"
                    
                    st.session_state.last_sync = sync_data
                    st.rerun()

            st.info("💡 Cliquez sur le bouton vert 'Enregistrer' pour déplacer votre texte vers la droite.")

        if audio_input:
            col_acc1, col_acc2 = st.columns(2)
            with col_acc1:
                if st.button("🚀 Transcrire & Traduire", use_container_width=True):
                    with st.spinner("Traitement..."):
                        new_txt = transcribe_audio_bytes(audio_input, src_lang)
                        new_trans = translate_text(new_txt, target_lang_name)
                        
                        # Accumulation sécurisée (Transcription)
                        if 'transcription' not in st.session_state or not st.session_state['transcription']:
                            st.session_state['transcription'] = new_txt
                        else:
                            st.session_state['transcription'] += f"\n\n---\n{new_txt}"
                        
                        # Accumulation sécurisée (Traduction)
                        if 'translation' not in st.session_state or not st.session_state['translation']:
                            st.session_state['translation'] = new_trans
                        else:
                            st.session_state['translation'] += f"\n\n---\n{new_trans}"
                        st.rerun()

    with col2:
        st.markdown("### 📝 2. Résultats")
        
        if 'transcription' in st.session_state and st.session_state['transcription']:
            original_text = st.text_area("📜 Historique des transcriptions", value=st.session_state['transcription'], height=250)
            st.session_state['transcription'] = original_text 

            if 'translation' in st.session_state:
                st.markdown(f"#### ✨ Traduction Globale ({target_lang_name})")
                
                # Affichage intelligent : Markdown pour les tableaux, zone de texte pour le reste
                if "|" in st.session_state['translation'] and "---" in st.session_state['translation']:
                    st.markdown(st.session_state['translation'])
                    with st.expander("📝 Modifier le texte brut / Tableau"):
                        current_translation = st.text_area("Éditeur Markdown", value=st.session_state['translation'], height=300)
                else:
                    current_translation = st.text_area("Texte", value=st.session_state['translation'], height=250)
                
                st.session_state['translation'] = current_translation
                # --------- LECTEUR AUDIO CLOUD-SAFE (GTTS MÉMOIRE) ---------
                st.markdown("---")
                
                # Mapper la langue cible vers le lecteur TTS (Google Base)
                lang_map = {
                    "Wolof": "fr",   # On lit en français la phonétique Wolof
                    "Français": "fr",
                    "Anglais": "en",
                    "Arabe": "ar",
                    "Espagnol": "es",
                    "Chinois": "zh-CN",
                    "Allemand": "de"
                }

                target_voice_lang = "fr" 
                for k, v in lang_map.items():
                    if k in target_lang_name:
                        target_voice_lang = v
                        break
                
                # Le nettoyage additionnel (les "cas désespérés" que l'IA oublierait)
                def clean_for_cloud_tts(text):
                    t = text
                    # Si c'est un tableau Markdown, on extrait uniquement la colonne "Texte à lire" pour le TTS
                    if "|" in t and "---" in t:
                        lines = t.split('\n')
                        extracted_lines = []
                        for line in lines:
                            if "|" in line and "---" not in line and "Texte à lire" not in line:
                                cols = [c.strip() for c in line.split('|') if c.strip()]
                                if len(cols) >= 1:
                                    extracted_lines.append(cols[0])
                        t = " ".join(extracted_lines)

                    t = t.replace("`", "'").replace("$", "")
                    if "Wolof" in target_lang_name:
                        # Remplacement des "w" par "ou" pour éviter la prononciation "V"
                        t = t.replace(" w", " ou").replace("W", "Ou").replace("wa", "oua").replace("wé", "oué")
                        t = t.replace("u", "ou").replace("U", "Ou")
                    return t
                
                safe_text = clean_for_cloud_tts(st.session_state['translation'])
                
                async def generate_edge_tts(text, voice):
                    communicate = edge_tts.Communicate(text, voice)
                    audio_data = b""
                    async for chunk in communicate.stream():
                        if chunk["type"] == "audio":
                            audio_data += chunk["data"]
                    return audio_data

                if st.button(f"🔊 Écouter avec {selected_voice_id.split('-')[-1].replace('Neural', '')} ({'Premium' if tts_engine == 'Google Cloud (Premium)' else 'Edge'})", use_container_width=True):
                    if safe_text.strip():
                        with st.spinner(f"Génération de l'audio..."):
                            try:
                                if tts_engine == "Google Cloud (Premium)":
                                    audio_content = generate_google_cloud_tts(safe_text, selected_voice_id, GEMINI_API_KEY)
                                elif tts_engine == "Gemini AI (Voix Native)":
                                    audio_content = generate_gemini_ai_audio(safe_text, target_lang_name)
                                else:
                                    audio_content = asyncio.run(generate_edge_tts(safe_text, selected_voice_id))
                                
                                if audio_content:
                                    st.audio(audio_content, format='audio/mp3' if tts_engine != "Gemini AI (Voix Native)" else 'audio/wav', autoplay=True)
                                else:
                                    st.warning("L'IA n'a pas pu générer l'audio.")
                            except Exception as e:
                                st.warning(f"Note : Passage au moteur de secours (gTTS)...")
                                try:
                                    tts = gTTS(text=safe_text, lang=target_voice_lang, slow=False)
                                    fp = io.BytesIO()
                                    tts.write_to_fp(fp)
                                    fp.seek(0)
                                    st.audio(fp, format='audio/mp3', autoplay=True)
                                except:
                                    st.error(f"Erreur audio : {e}")
                    else:
                        st.warning("Rien à lire.")
                # (Ancien lecteur Javascript supprimé, remplacé par st.audio)

            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button(f"🌍 Traduire tout en {target_lang_name}", use_container_width=True):
                    with st.spinner(f"Traduction de l'historique..."):
                        translation = translate_text(original_text, target_lang_name)
                        st.session_state['translation'] = translation
                        st.rerun()

            with col_btn2:
                # Créer un contenu TXT combiné
                full_export = f"=== TRANSCRIPTION ({src_lang}) ===\n\n{original_text}"
                if 'translation' in st.session_state:
                    full_export += f"\n\n\n=== TRADUCTION ({target_lang_name}) ===\n\n{st.session_state['translation']}"
                
                st.download_button("📥 Exporter Tout (Transcription + Traduction)", full_export, "session_wolof_ai.txt", use_container_width=True)

            if st.button("🗑️ Effacer l'historique", type="secondary", use_container_width=True):
                del st.session_state['transcription']
                if 'translation' in st.session_state: del st.session_state['translation']
                st.rerun()
        
        elif selected_menu == "📄 Document" and 'doc_translation' in st.session_state:
            st.markdown("### 📝 Traduction du Document")
            
            # Affichage intelligent : Markdown pour les tableaux, zone de texte pour le reste
            if "|" in st.session_state['doc_translation'] and "---" in st.session_state['doc_translation']:
                st.markdown(st.session_state['doc_translation'])
                with st.expander("📝 Modifier le texte brut / Tableau"):
                    edited_translation = st.text_area("Éditeur Markdown", value=st.session_state['doc_translation'], height=400)
            else:
                edited_translation = st.text_area(
                    "Vous pouvez modifier la traduction ici avant de télécharger :",
                    st.session_state['doc_translation'],
                    height=400,
                    key="doc_edit_area"
                )
            
            st.session_state['doc_translation'] = edited_translation
            
            st.divider()
            st.markdown("#### 📥 Téléchargement")
            
            col_dl1, col_dl2 = st.columns(2)
            
            # Génération DOCX
            with col_dl1:
                doc_obj = Document()
                doc_obj.add_heading('Traduction XAMLE IA STUDIO', 0)
                doc_obj.add_paragraph(edited_translation)
                
                buf_docx = io.BytesIO()
                doc_obj.save(buf_docx)
                st.download_button(
                    label="📥 Télécharger en DOCX",
                    data=buf_docx.getvalue(),
                    file_name=f"traduction_xamle_{st.session_state.get('last_doc_name', 'doc').replace('.pdf','')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            # Génération PDF
            with col_dl2:
                buf_pdf = io.BytesIO()
                c = canvas.Canvas(buf_pdf, pagesize=letter)
                width, height = letter
                
                c.setFont("Helvetica-Bold", 16)
                c.drawString(50, height - 50, "Traduction XAMLE IA STUDIO")
                
                c.setFont("Helvetica", 12)
                text_obj = c.beginText(50, height - 80)
                # Gestion simple du Wrapping
                lines = edited_translation.split('\n')
                for line in lines:
                    # Tronquer les lignes trop longues pour le canvas de base (approche simplifiée)
                    # Note: ReportLab nécessite un traitement plus fin pour le wrap auto
                    text_obj.textLine(line[:100]) 
                
                c.drawText(text_obj)
                c.showPage()
                c.save()
                
                st.download_button(
                    label="📥 Télécharger en PDF",
                    data=buf_pdf.getvalue(),
                    file_name=f"traduction_xamle_{st.session_state.get('last_doc_name', 'doc').replace('.pdf','')}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )

        else:
            st.info("Les transcriptions s'ajouteront ici progressivement.")
            st.image("https://cdni.iconscout.com/illustration/premium/thumb/empty-state-2130362-1800926.png", use_container_width=True)

if __name__ == "__main__":
    main()
